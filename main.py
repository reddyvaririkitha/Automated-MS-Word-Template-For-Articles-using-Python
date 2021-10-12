import docx
from docx.oxml.xmlchemy import OxmlElement
from docx.oxml.shared import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


document = docx.Document()

sec_pr = document.sections[0]._sectPr  # get the section properties

# creating new borders
pg_borders = OxmlElement('w:pgBorders')

# relative positioning of the borders should be calculated
pg_borders.set(qn('w:offsetFrom'), 'page')
for border_name in ('top', 'left', 'bottom', 'right',):  # set all borders
    border_el = OxmlElement(f'w:{border_name}')
    border_el.set(qn('w:val'), 'single')  # a single line
    border_el.set(qn('w:sz'), '4')  # for meaning of  remaining attrs please look docs
    border_el.set(qn('w:space'), '24')
    border_el.set(qn('w:color'), 'auto')
    pg_borders.append(border_el)  # register single border to border

# applying border changes to section
sec_pr.append(pg_borders)

# Defining heading
Heading = document.add_heading()

Heading = Heading.add_run("Prevention by Disinfection:")

# Heading styles- font
font = Heading.font
font.name = 'Calibri Light'
# Heading styles- size
font.size = Pt(60)
# Heading styles- colour
font.color.rgb = RGBColor(127, 209, 59)


# here 'Paragraph' is an object
Paragraph = document.add_paragraph()

# Paragraph styles- alignment
paragraph_format = Paragraph.paragraph_format
paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Paragraph styles- font
paragraph_style = document.styles['Normal']
paragraph_font = paragraph_style.font
paragraph_font.name = 'Calibri'

# Paragraph styles- size
paragraph_font.size = Pt(12)

# Paragraph styles- colour
paragraph_font.color.rgb = RGBColor(0, 0, 0)

Paragraph.add_run("\nThis is the word document generated using python. This is a basic template and many other things "
                  "can be added to this template such as- pictures, tables, bullets, header and footers, etc.")

# saving the document
document.save('template.docx')
